import docx
from googletrans import Translator


def translate_text(text, translator, target_language):
    translated_text = translator.translate(text, dest=target_language).text
    return translated_text


def translate_word_file(input_file, output_file, target_language):
    # Read the content of the Word file
    try:
        doc = docx.Document(input_file)
    except Exception as e:
        print(f"Error reading the Word file: {e}")
        return

    # Translate the content
    translator = Translator()
    translated_doc = docx.Document()

    for paragraph in doc.paragraphs:
        translated_paragraph = translated_doc.add_paragraph()
        for run in paragraph.runs:
            translated_text = translate_text(
                run.text, translator, target_language)
            translated_run = translated_paragraph.add_run(translated_text)
            # Copy over font properties
            translated_run.bold = run.bold
            translated_run.italic = run.italic
            translated_run.underline = run.underline
            translated_run.font.name = run.font.name
            translated_run.font.size = run.font.size
            translated_run.font.color.rgb = run.font.color.rgb

        # Copy over paragraph style
        translated_paragraph.style = paragraph.style

    # Save the translated document
    try:
        translated_doc.save(output_file)
        print(f"Translation saved to {output_file}")
    except Exception as e:
        print(f"Error saving the translated content: {e}")


if __name__ == "__main__":
    # Get input parameters from the command line
    input_file = input("Enter the path to the Word file: ").strip()
    output_file = input(
        "Enter the path to save the translated Word file (include .docx extension): ").strip()
    target_language = input(
        "Enter the target language (e.g., 'fr' for French): ").strip()

    # Call the translation function
    translate_word_file(input_file, output_file, target_language)
